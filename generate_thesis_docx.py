#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成论文Word文档脚本
按照国家开放大学学士学位论文格式要求生成论文初版
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    
    if level == 1:
        # 一级标题：黑体三号加粗居中
        set_chinese_font(run, '黑体', 16, True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_before = Pt(18)
        para.space_after = Pt(12)
    elif level == 2:
        # 二级标题：黑体四号加粗左对齐
        set_chinese_font(run, '黑体', 14, True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.space_before = Pt(12)
        para.space_after = Pt(6)
    elif level == 3:
        # 三级标题：黑体小四号加粗左对齐
        set_chinese_font(run, '黑体', 12, True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.space_before = Pt(6)
        para.space_after = Pt(6)
    
    return para

def add_normal_text(doc, text):
    """添加正文段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '宋体', 12, False)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = Cm(0.75)  # 首行缩进2字符
    para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    return para

def read_chapter_file(filepath):
    """读取章节文件内容"""
    if not os.path.exists(filepath):
        return ""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def process_chapter_content(doc, content):
    """处理章节内容，识别标题和正文"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 跳过分隔线
        if line.startswith('='):
            continue
        
        # 识别一级标题（如：一、综述）
        if line.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、')):
            add_heading_custom(doc, line, level=1)
        # 识别二级标题（如：（一）系统建设背景）
        elif line.startswith('（') and '）' in line[:5]:
            add_heading_custom(doc, line, level=2)
        # 识别三级标题（以数字开头，如：1. LLM）
        elif line[0].isdigit() and '. ' in line[:5]:
            add_heading_custom(doc, line, level=3)
        # 普通正文
        else:
            if line:
                add_normal_text(doc, line)

def main():
    """主函数"""
    print("开始生成论文Word文档...")
    
    # 创建文档对象
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4纸高度
    section.page_width = Cm(21)     # A4纸宽度
    section.left_margin = Cm(3)     # 左边距3cm
    section.right_margin = Cm(2)    # 右边距2cm
    section.top_margin = Cm(2.5)    # 上边距2.5cm
    section.bottom_margin = Cm(2.5) # 下边距2.5cm
    
    # 章节文件列表
    chapters = [
        ('论文第一章综述.TXT', '第一章 综述'),
        ('论文第二章需求分析.TXT', '第二章 需求分析'),
        ('论文第三章系统总体设计.TXT', '第三章 系统总体设计'),
        ('论文第四章系统详细设计与实现.TXT', '第四章 系统详细设计与实现'),
        ('论文第五章系统测试.TXT', '第五章 系统测试'),
        ('论文第六章系统部署与运维.TXT', '第六章 系统部署与运维'),
        ('论文第七章结论.TXT', '第七章 结论'),
        ('论文致谢.TXT', '致谢'),
        ('论文参考文献.TXT', '参考文献'),
    ]
    
    # 处理每个章节
    for filename, chapter_title in chapters:
        filepath = os.path.join(os.path.dirname(__file__), filename)
        print(f"处理章节: {chapter_title}")
        
        # 添加章节标题（作为一级标题）
        if not chapter_title.startswith(('致谢', '参考文献')):
            add_heading_custom(doc, chapter_title, level=1)
        
        # 读取并处理章节内容
        content = read_chapter_file(filepath)
        if content:
            process_chapter_content(doc, content)
        else:
            print(f"警告: 文件 {filename} 不存在或为空")
        
        # 章节之间添加分页符（除最后一章）
        if filename != chapters[-1][0]:
            doc.add_page_break()
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(__file__), '论文初版.docx')
    doc.save(output_path)
    print(f"\n论文Word文档已生成: {output_path}")
    print("文件名: 论文初版.docx")
    print("请打开文档检查格式，如有需要可手动调整")

if __name__ == '__main__':
    main()
