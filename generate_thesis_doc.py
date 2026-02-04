#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成符合学位论文格式的Word文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', font_size=12):
    """设置中文字体（中文：宋体/黑体，英文数字：Times New Roman）"""
    run.font.name = 'Times New Roman'  # 英文和数字字体
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体

def add_heading_custom(doc, text, level=1):
    """添加自定义格式的标题"""
    # 一级标题（章）前添加分页符
    if level == 1:
        doc.add_page_break()
    
    para = doc.add_paragraph()
    # 顶级标题和一级标题都居中
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if level in [0, 1] else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    
    if level == 0:  # 摘要、Abstract等顶级标题
        set_chinese_font(run, '黑体', 16)
        run.bold = True
    elif level == 1:  # 一级标题（章）
        set_chinese_font(run, '黑体', 16)
        run.bold = True
    elif level == 2:  # 二级标题
        set_chinese_font(run, '黑体', 14)
        run.bold = True
    elif level == 3:  # 三级标题
        set_chinese_font(run, '黑体', 12)
        run.bold = True
    
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_normal_paragraph(doc, text):
    """添加正文段落"""
    if not text or text.strip() == '':
        return
    
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '宋体', 12)
    
    para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def process_line(doc, line):
    """处理单行文本"""
    line = line.strip()
    
    if not line or line == '=' * len(line):
        return
    
    # 判断标题级别
    if line.startswith('摘  要') or line.startswith('Abstract') or line.startswith('参考文献') or line.startswith('致  谢') or line.startswith('附  录'):
        add_heading_custom(doc, line, level=0)
    elif line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or line.startswith('四、') or line.startswith('五、') or line.startswith('六、'):
        add_heading_custom(doc, line, level=1)
    elif line.startswith('（一）') or line.startswith('（二）') or line.startswith('（三）') or line.startswith('（四）') or line.startswith('（五）') or line.startswith('（六）'):
        add_heading_custom(doc, line, level=2)
    elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.') or line.startswith('5.'):
        if len(line) < 30:  # 短行作为三级标题
            add_heading_custom(doc, line, level=3)
        else:
            add_normal_paragraph(doc, line)
    elif line.startswith('关键词：') or line.startswith('Keywords:'):
        para = doc.add_paragraph()
        run = para.add_run(line)
        set_chinese_font(run, '宋体', 12)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif line.startswith('附录A：') or line.startswith('附录B：') or line.startswith('附录C：'):
        add_heading_custom(doc, line, level=2)
    elif line.startswith('表') and '-' in line:  # 表格标题
        para = doc.add_paragraph()
        run = para.add_run(line)
        set_chinese_font(run, '宋体', 10.5)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_normal_paragraph(doc, line)

def add_cover_page(doc):
    """添加封面"""
    # 标题：国家开放大学
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(72)
    run = para.add_run('国家开放大学')
    set_chinese_font(run, '黑体', 22)
    run.bold = True
    
    # 学士学位论文
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    run = para.add_run('学士学位论文')
    set_chinese_font(run, '黑体', 18)
    run.bold = True
    
    # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 题目
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    run = para.add_run('题目：企业大语言模型训练管理平台的设计与实现')
    set_chinese_font(run, '黑体', 16)
    
    # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 信息部分（居中）
    info_items = [
        '分部：',
        '学习中心：',
        '专业：计算机科学与技术',
        '入学时间：',
        '学号：',
        '姓名：',
        '指导教师：',
        '论文完成日期：    年    月'
    ]
    
    for item in info_items:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(item)
        set_chinese_font(run, '宋体', 14)
    
    # 分页符
    doc.add_page_break()

def add_declaration(doc):
    """添加原创性声明和授权声明"""
    # 原创性声明标题
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    run = para.add_run('学位论文原创性声明')
    set_chinese_font(run, '黑体', 16)
    run.bold = True
    
    # 原创性声明内容
    declaration_text = '本人郑重声明：所呈交的学位论文，是本人在导师指导下，进行研究工作所取得的成果。除文中已经注明引用的内容外，本学位论文的研究成果不包含任何他人创作的、已公开发表或者没有公开发表的作品的内容。对本论文所涉及的研究工作做出贡献的其他个人和集体，均已在文中以明确方式标明。本学位论文原创性声明的法律责任由本人承担。'
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(declaration_text)
    set_chinese_font(run, '宋体', 12)
    
    # 签名
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    run = para.add_run('作者签名：              日期：    ')
    set_chinese_font(run, '宋体', 12)
    
    # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 授权声明标题
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    run = para.add_run('学位论文版权使用授权声明')
    set_chinese_font(run, '黑体', 16)
    run.bold = True
    
    # 授权声明内容
    authorization_text = '本人完全了解国家开放大学关于收集、保存、使用学位论文的规定，同意如下各项内容：按照学校要求提交学位论文的印刷本和电子版本；学校有权保存学位论文的印刷本和电子版，并采用影印、缩印、扫描、数字化或其他手段保存论文；学校有权提供目录检索以及提供本学位论文全文或者部分的阅览服务，以及出版学位论文；学校有权按有关规定向国家有关部门或者机构送交论文的复印件和电子版；在不以营利为目的的前提下，学校可以适当复制论文的部分或全部内容用于学术活动。'
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(authorization_text)
    set_chinese_font(run, '宋体', 12)
    
    # 签名
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    run = para.add_run('作者签名：              日期：    ')
    set_chinese_font(run, '宋体', 12)
    
    # 分页符
    doc.add_page_break()

def add_toc(doc):
    """添加目录（手动生成）"""
    # 目录标题
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    run = para.add_run('目  录')
    set_chinese_font(run, '黑体', 16)
    run.bold = True
    
    doc.add_paragraph()
    
    # 目录项（根据论文目录结构(新).txt）
    toc_items = [
        ('摘  要', 'Ⅰ'),
        ('Abstract', 'Ⅱ'),
        ('一、综述', '1'),
        ('    （一）系统建设背景', '1'),
        ('    （二）研究意义', '2'),
        ('    （三）术语定义', '3'),
        ('    （四）技术选型', '3'),
        ('    （五）系统运行环境', '6'),
        ('二、需求分析', '8'),
        ('    （一）系统业务总体需求', '8'),
        ('    （二）系统功能需求分析', '9'),
        ('    （三）系统非功能需求分析', '17'),
        ('三、系统总体设计', '19'),
        ('    （一）业务流程设计', '19'),
        ('    （二）系统架构设计', '23'),
        ('    （三）系统网络拓扑设计', '26'),
        ('    （四）数据库设计', '27'),
        ('四、系统详细设计与实现', '34'),
        ('    （一）用户认证模块', '34'),
        ('    （二）模型配置管理模块', '41'),
        ('    （三）模型对话模块', '48'),
        ('    （四）模型对比测试模块', '57'),
        ('    （五）模型训练管理模块', '64'),
        ('    （六）训练可视化监控模块', '71'),
        ('    （七）系统提示词管理模块', '78'),
        ('    （八）系统管理模块', '84'),
        ('五、系统测试', '91'),
        ('    （一）测试目的', '91'),
        ('    （二）测试环境', '91'),
        ('    （三）功能测试', '92'),
        ('    （四）性能测试', '98'),
        ('    （五）安全性测试', '101'),
        ('    （六）测试结果分析', '104'),
        ('六、总结与展望', '105'),
        ('    （一）工作总结', '105'),
        ('    （二）创新点', '106'),
        ('    （三）存在的问题', '107'),
        ('    （四）未来展望', '108'),
        ('参考文献', '109'),
        ('致  谢', '110'),
        ('附  录', '111'),
        ('    附录A：系统主要界面截图', '111'),
        ('    附录B：核心代码清单', '115'),
        ('    附录C：数据库表结构详细说明', '120'),
    ]
    
    for item, page in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        
        # 添加目录项文本
        run = para.add_run(item)
        set_chinese_font(run, '宋体', 12)
        
        # 添加制表符和页码
        run = para.add_run('\t' + page)
        set_chinese_font(run, '宋体', 12)
    
    # 分页符
    doc.add_page_break()

def main():
    """主函数"""
    # 创建文档
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4纸高度
    section.page_width = Cm(21)     # A4纸宽度
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # 添加封面
    add_cover_page(doc)
    
    # 添加原创性声明和授权声明
    add_declaration(doc)
    
    # 添加目录
    add_toc(doc)
    
    # 读取论文内容
    with open('论文第三版.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按行处理
    lines = content.split('\n')
    for line in lines:
        process_line(doc, line)
    
    # 保存文档
    doc.save('论文第三版.docx')
    print('论文第三版.docx 已生成完成！')

if __name__ == '__main__':
    main()
